from pathlib import Path
from datetime import date
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/rogergv/Documents/SoftvibesLab/BeGlobal")
OUTPUT = ROOT / "contrtos" / "CONTRATO_PRESTACION_SERVICIOS_SISTEMA_BE_GLOBAL_PRO_SOFTVIBES_3_AGENTES.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "FFF4CC"
WHITE = "FFFFFF"
BLACK = "000000"
RISK = "9B1C1C"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="C8CDD3", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Column widths must total {TABLE_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_keep(paragraph, next_=False, together=True):
    paragraph.paragraph_format.keep_together = together
    paragraph.paragraph_format.keep_with_next = next_


def shade_paragraph(paragraph, fill, border=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        for edge in ("top", "left", "bottom", "right"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "6")
            node.set(qn("w:space"), "4")
            node.set(qn("w:color"), border)
            p_bdr.append(node)


def add_page_field(paragraph, field_name):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 14, 8),
        ("Heading 2", 13, BLUE, 11, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for custom_name, base, size, color, bold, italic, before, after in (
        ("Contract Title", "Normal", 22, NAVY, True, False, 0, 6),
        ("Contract Subtitle", "Normal", 13, MUTED, False, False, 0, 18),
        ("Clause Body", "Normal", 11, BLACK, False, False, 0, 6),
        ("Definition", "Normal", 10.5, BLACK, False, False, 0, 5),
        ("Small Note", "Normal", 9, MUTED, False, False, 0, 4),
        ("Signature Label", "Normal", 9, MUTED, True, False, 0, 2),
    ):
        if custom_name in styles:
            style = styles[custom_name]
        else:
            style = styles.add_style(custom_name, 1)
            style.base_style = styles[base]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
        style.font.italic = italic
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25 if custom_name in ("Clause Body", "Definition") else 1.15


def add_numbering_definition(doc, num_fmt, text, left=540, hanging=270, font=None):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs, default=0) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.extend([tabs, ind])
    lvl.extend([start, fmt, lvl_text, suff, p_pr])
    if font:
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), font)
        r_fonts.set(qn("w:hAnsi"), font)
        r_pr.append(r_fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


def add_list_item(doc, text, num_id, bold_prefix=None, after=4):
    p = doc.add_paragraph(style="Clause Body")
    apply_numbering(p, num_id)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        append_text_with_placeholders(p, text[len(bold_prefix):])
    else:
        append_text_with_placeholders(p, text)
    set_keep(p)
    return p


def append_text_with_placeholders(paragraph, text, size=11):
    """Add text and highlight bracketed fields that still require completion."""
    for part in re.split(r"(\[[^\]]+\])", text):
        if not part:
            continue
        if part.startswith("[") and part.endswith("]"):
            run = add_placeholder_run(paragraph, part)
            run.font.size = Pt(size)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def add_body(doc, text, bold_prefix=None, italic=False, align=None, keep=False):
    p = doc.add_paragraph(style="Clause Body")
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    set_keep(p, together=keep)
    return p


def add_definition(doc, term, definition):
    p = doc.add_paragraph(style="Definition")
    r1 = p.add_run(f"{term}. ")
    set_run_font(r1, size=10.5, color=DARK_BLUE, bold=True)
    r2 = p.add_run(definition)
    set_run_font(r2, size=10.5)
    set_keep(p)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep(p, next_=True)
    return p


def add_placeholder_run(paragraph, text):
    run = paragraph.add_run(text)
    set_run_font(run, color=RISK, bold=True)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return run


def fill_cell(cell, text, bold=False, color=BLACK, size=9.5, placeholder=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.clear()
    if placeholder:
        run = add_placeholder_run(p, text)
        run.font.size = Pt(size)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_kv_table(doc, rows, widths=(2700, 6660), header=None):
    count = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=count, cols=2)
    set_table_geometry(table, list(widths))
    row_idx = 0
    if header:
        fill_cell(table.cell(0, 0), header[0], bold=True, color=WHITE)
        fill_cell(table.cell(0, 1), header[1], bold=True, color=WHITE)
        set_cell_shading(table.cell(0, 0), NAVY)
        set_cell_shading(table.cell(0, 1), NAVY)
        set_repeat_table_header(table.rows[0])
        row_idx = 1
    for idx, (label, value, is_placeholder) in enumerate(rows):
        fill_cell(table.cell(row_idx + idx, 0), label, bold=True, color=DARK_BLUE)
        set_cell_shading(table.cell(row_idx + idx, 0), LIGHT_GRAY)
        fill_cell(table.cell(row_idx + idx, 1), value, placeholder=is_placeholder)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_data_table(doc, headers, rows, widths, sizes=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for col, header in enumerate(headers):
        fill_cell(table.cell(0, col), header, bold=True, color=WHITE, size=9)
        set_cell_shading(table.cell(0, col), NAVY)
    for row_idx, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            size = sizes[col] if sizes else 9
            fill_cell(table.cell(row_idx, col), value, size=size)
            if row_idx % 2 == 0:
                set_cell_shading(table.cell(row_idx, col), "FAFBFC")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_signature_table(doc, parties):
    table = doc.add_table(rows=1, cols=2)
    # Signature blocks intentionally use zero indent so both columns stay inside
    # the printable width in Pages/Word; the cell padding provides text inset.
    set_table_geometry(table, [4680, 4680], indent_dxa=0)
    for col, (party, representative) in enumerate(parties):
        cell = table.cell(0, col)
        fill_cell(cell, f"POR {party}", bold=True, color=NAVY, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("________________________________")
        set_run_font(r, size=10, color=MUTED)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_placeholder_run(p, representative)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Nombre y firma del representante")
        set_run_font(r, size=8.5, color=MUTED)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_placeholder_run(p, "[CARGO]")
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_placeholder_run(p, "[FECHA]")
    return table


def configure_page(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def set_header_footer(section):
    for header in (section.header, section.even_page_header):
        p = header.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run("BORRADOR PARA REVISIÓN LEGAL")
        set_run_font(r1, size=8, color=RISK, bold=True)
        r2 = p.add_run("  |  BeGlobal - Softvibes")
        set_run_font(r2, size=8, color=MUTED)

    for footer in (section.footer, section.even_page_footer):
        p = footer.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run("Página ")
        set_run_font(r, size=8.5, color=MUTED)
        add_page_field(p, "PAGE")
        r = p.add_run(" de ")
        set_run_font(r, size=8.5, color=MUTED)
        add_page_field(p, "NUMPAGES")


def add_cover(doc, bullet_num):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("BORRADOR PARA REVISIÓN LEGAL")
    set_run_font(r, size=10, color=RISK, bold=True)

    p = doc.add_paragraph(style="Contract Title")
    r = p.add_run("CONTRATO MARCO DE PRESTACIÓN DE SERVICIOS")
    set_run_font(r, size=22, color=NAVY, bold=True)

    p = doc.add_paragraph(style="Contract Subtitle")
    r = p.add_run("Implementación, personalización, activación y soporte de Agentes IA Ecommerce")
    set_run_font(r, size=13, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r1 = p.add_run("Entre ")
    set_run_font(r1, size=12, color=MUTED)
    r2 = p.add_run("SISTEMA BE GLOBAL PRO, S.A. DE C.V.")
    set_run_font(r2, size=12, color=NAVY, bold=True)
    r3 = p.add_run(" y ")
    set_run_font(r3, size=12, color=MUTED)
    r4 = p.add_run("SOFTVIBES")
    set_run_font(r4, size=12, color=NAVY, bold=True)

    add_kv_table(
        doc,
        [
            ("Fecha del borrador", "31 de julio de 2026", False),
            ("Versión", "0.3 - tres Agentes y precio con IVA incluido", False),
            ("Jurisdicción propuesta", "Hermosillo, Sonora, México", False),
            ("Estado", "NO FIRMAR HASTA COMPLETAR Y VALIDAR LOS CAMPOS RESALTADOS", False),
        ],
        widths=(2700, 6660),
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.line_spacing = 1.15
    shade_paragraph(p, CALLOUT, border="D6B656")
    r1 = p.add_run("AVISO DE PREPARACIÓN. ")
    set_run_font(r1, size=9.5, color=RISK, bold=True)
    r2 = p.add_run(
        "Este documento es un borrador contractual elaborado a partir del contrato de referencia y del alcance comercial documentado en el proyecto BeGlobal. Los datos fiscales de BEGLOBAL se incorporaron de su Constancia de Situación Fiscal emitida el 12 de junio de 2026. Requiere revisión de asesoría jurídica mexicana y confirmación corporativa, de representación y operativa antes de firma."
    )
    set_run_font(r2, size=9.5, color=BLACK)

    add_heading(doc, "Campos que deben completarse antes de firma", 2)
    for item in (
        "Nombre, cargo, instrumento de poder, identificación y correo del representante legal de BEGLOBAL.",
        "Razón social o nombre fiscal, RFC y domicilio del Prestador.",
        "Facultades, cargo, identificación y correo del representante del Prestador.",
        "Canal de soporte, horario, costos recurrentes y proveedores tecnológicos autorizados.",
        "Calendario definitivo del piloto y esquema de facturación.",
    ):
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "Criterios jurídicos incorporados", 2)
    add_body(
        doc,
        "El borrador considera el régimen mexicano de prestación de servicios y buena fe contractual; la regulación vigente sobre obras por encargo y derechos patrimoniales; las obligaciones de protección de datos personales; la validez de mensajes de datos y firma electrónica; y la necesidad de mantener una relación realmente independiente, sin subordinación laboral."
    )
    p = doc.add_paragraph(style="Small Note")
    r = p.add_run(
        "Base normativa revisada: Código Civil Federal; Código de Comercio; Ley Federal del Derecho de Autor; Ley Federal de Protección de Datos Personales en Posesión de los Particulares; y Ley Federal del Trabajo, textos vigentes consultados al 31 de julio de 2026."
    )
    set_run_font(r, size=9, color=MUTED, italic=True)
    doc.add_page_break()


def build_document():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = True
    configure_styles(doc)
    configure_page(doc)
    for section in doc.sections:
        set_header_footer(section)

    bullet_num = add_numbering_definition(doc, "bullet", "•", left=540, hanging=270, font="Symbol")
    decimal_num = add_numbering_definition(doc, "decimal", "%1.", left=540, hanging=270)

    add_cover(doc, bullet_num)

    p = doc.add_paragraph(style="Contract Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONTRATO MARCO DE PRESTACIÓN DE SERVICIOS")
    set_run_font(r, size=18, color=NAVY, bold=True)
    p = doc.add_paragraph(style="Contract Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Agentes IA Ecommerce para BeGlobal")
    set_run_font(r, size=12.5, color=MUTED)

    p = doc.add_paragraph(style="Clause Body")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("QUE CELEBRAN, POR UNA PARTE, ")
    set_run_font(r, bold=True)
    r = p.add_run("SISTEMA BE GLOBAL PRO, SOCIEDAD ANÓNIMA DE CAPITAL VARIABLE")
    set_run_font(r, bold=True)
    r = p.add_run(", con RFC SBG220718U55, titular o licenciataria de la marca comercial “BE GLOBAL PRO”, representada en este acto por ")
    set_run_font(r)
    add_placeholder_run(p, "[NOMBRE COMPLETO DEL REPRESENTANTE LEGAL]")
    r = p.add_run(", a quien en lo sucesivo se denominará “BEGLOBAL”; Y, POR LA OTRA, ")
    set_run_font(r)
    add_placeholder_run(p, "[NOMBRE O RAZÓN SOCIAL DEL PRESTADOR]")
    r = p.add_run(", que opera comercialmente como “SOFTVIBES”, representada por ")
    set_run_font(r)
    add_placeholder_run(p, "[JOSÉ ROGELIO GARCÍA VITAL / REPRESENTANTE AUTORIZADO]")
    r = p.add_run(", a quien se denominará el “PRESTADOR”; conjuntamente, las “PARTES”, al tenor de las siguientes declaraciones y cláusulas.")
    set_run_font(r)

    add_heading(doc, "DECLARACIONES", 1)
    add_heading(doc, "I. Declara BEGLOBAL", 2)
    for item in (
        "Que es una sociedad anónima de capital variable denominada SISTEMA BE GLOBAL PRO, con RFC SBG220718U55 y estatus activo en el padrón federal de contribuyentes conforme a la Constancia de Situación Fiscal emitida el 12 de junio de 2026.",
        "Que tiene su domicilio fiscal en Boulevard Paseo de las Quintas número 184, colonia Portal del Lago, código postal 83240, Hermosillo, Sonora, México.",
        "Que comparece por conducto de [NOMBRE COMPLETO DEL REPRESENTANTE LEGAL], en su carácter de [CARGO], quien manifiesta contar con facultades suficientes para obligarla en los términos del presente Contrato, mismas que no le han sido revocadas ni limitadas.",
        "Que desarrolla su actividad bajo el nombre comercial SISTEMA BE GLOBAL PRO, utiliza la marca BE GLOBAL PRO y desarrolla programas, contenidos, metodologías y servicios relacionados con ecommerce, formación, acompañamiento y operación digital.",
        "Que desea incorporar Agentes IA Ecommerce como apoyo a su equipo y a determinados clientes o miembros, sin sustituir la supervisión humana ni garantizar resultados comerciales.",
        "Que cuenta o contará con las autorizaciones necesarias sobre los materiales, marcas, datos y contenidos que proporcione al PRESTADOR.",
    ):
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "II. Declara el PRESTADOR", 2)
    for item in (
        "Que es una persona [MORAL/FÍSICA CON ACTIVIDAD EMPRESARIAL] con RFC [RFC], domicilio en [DOMICILIO COMPLETO] y capacidad técnica y jurídica para prestar los Servicios.",
        "Que dispone de conocimientos en software, automatización, inteligencia artificial, agentes digitales, ecommerce, contenidos y operación tecnológica.",
        "Que prestará los Servicios con autonomía técnica y administrativa, mediante recursos propios o terceros autorizados, sin subordinación respecto de BEGLOBAL.",
        "Que informará a BEGLOBAL sobre dependencias, limitaciones y costos de terceros relevantes antes de asumir compromisos que generen cargos adicionales.",
    ):
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "III. Declaran ambas PARTES", 2)
    for item in (
        "Que se reconocen capacidad y representación suficientes para obligarse.",
        "Que su consentimiento se encuentra libre de error, dolo, violencia o mala fe.",
        "Que el presente Contrato es de naturaleza civil o mercantil, según corresponda a la actividad de las PARTES, y se ejecutará conforme a la buena fe, la colaboración y la trazabilidad documental.",
    ):
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "DEFINICIONES", 1)
    definitions = (
        ("Agente IA Ecommerce", "Configuración de software y modelos de inteligencia artificial que recibe instrucciones y contenidos para orientar, generar borradores o apoyar tareas de ecommerce dentro del alcance autorizado."),
        ("Activación", "Momento en que un Agente ha superado la revisión técnica mínima, se asigna a un cliente final o usuario autorizado y queda disponible en el canal pactado."),
        ("Cliente Final", "Miembro, socio, alumno o cliente de BEGLOBAL para quien se personalice o habilite un Agente."),
        ("Entregables", "Resultados expresamente descritos en una Orden de Trabajo, incluidos configuraciones, documentos, bases de conocimiento, reportes o accesos."),
        ("Materiales de BEGLOBAL", "Marcas, metodologías, textos, imágenes, videos, bases de conocimiento, datos, credenciales temporales e instrucciones proporcionados por BEGLOBAL."),
        ("Orden de Trabajo", "Anexo o documento firmado por las PARTES que define alcance, volumen, precio, calendario, criterios de aceptación, soporte y dependencias de un servicio específico."),
        ("Servicios de Terceros", "Modelos de IA, alojamiento, mensajería, APIs, tiendas, plataformas, licencias y otros servicios no controlados por el PRESTADOR."),
    )
    for term, definition in definitions:
        add_definition(doc, term, definition)

    clauses = []

    add_heading(doc, "PRIMERA. OBJETO", 1)
    add_body(
        doc,
        "El PRESTADOR se obliga a diseñar, configurar, personalizar, probar, activar y, cuando se pacte, soportar Agentes IA Ecommerce y componentes relacionados para BEGLOBAL, conforme a este Contrato y a cada Orden de Trabajo. BEGLOBAL se obliga a proporcionar insumos, aprobaciones y pagos en los términos acordados."
    )
    add_body(
        doc,
        "Los Servicios son obligaciones de medios y diligencia profesional. Ningún Entregable implica garantía de ventas, ingresos, posicionamiento, aprobación de plataformas, disponibilidad ininterrumpida de terceros ni resultados empresariales específicos."
    )

    add_heading(doc, "SEGUNDA. ÓRDENES DE TRABAJO Y PRELACIÓN", 1)
    add_body(doc, "Cada proyecto, lote o activación podrá documentarse en una Orden de Trabajo que deberá indicar, como mínimo:")
    for item in (
        "alcance incluido y exclusiones;",
        "Entregables y criterios de aceptación;",
        "volumen de Agentes o usuarios;",
        "precio, impuestos, costos de terceros y forma de pago;",
        "calendario, dependencias y responsables;",
        "soporte, niveles de atención y periodo de garantía; y",
        "categorías de datos y proveedores tecnológicos autorizados.",
    ):
        add_list_item(doc, item, bullet_num)
    add_body(
        doc,
        "En caso de contradicción, prevalecerán: (i) el convenio modificatorio posterior; (ii) el anexo de datos y seguridad para asuntos de datos personales; (iii) la Orden de Trabajo para su alcance comercial y técnico; y (iv) este Contrato. Ninguna conversación informal modificará el alcance sin confirmación escrita de representantes autorizados."
    )

    add_heading(doc, "TERCERA. ALCANCE BASE", 1)
    add_body(doc, "Salvo que una Orden de Trabajo disponga otra cosa, el servicio base por Agente podrá incluir:")
    for item in (
        "personalización con nombre o identidad comercial del Cliente Final;",
        "diagnóstico inicial de etapa, necesidades y siguiente paso;",
        "guía sobre ecommerce, contenido, ofertas, ventas y canales digitales;",
        "carga o vinculación de una base inicial de conocimiento aprobada por BEGLOBAL;",
        "pruebas funcionales, ajustes iniciales y Activación;",
        "entrega de instrucciones de uso y canal de soporte correctivo durante el periodo pactado.",
    ):
        add_list_item(doc, item, bullet_num)
    add_body(
        doc,
        "No se incluyen integraciones avanzadas, CRM a medida, automatizaciones complejas, carga masiva de catálogo, administración integral del negocio, atención humana ilimitada, desarrollo no descrito, costos de terceros ni ejecución autónoma de transacciones. Tales conceptos requerirán una Orden de Trabajo adicional."
    )

    add_heading(doc, "CUARTA. CONTROL DE CAMBIOS", 1)
    add_body(
        doc,
        "Cualquier modificación de alcance, volumen, calendario, entregables, seguridad o costos deberá documentarse mediante solicitud de cambio. El PRESTADOR informará el impacto estimado y no estará obligado a iniciar el cambio hasta recibir aprobación escrita de BEGLOBAL. Los retrasos atribuibles a insumos, accesos o aprobaciones de BEGLOBAL ajustarán razonablemente el calendario."
    )

    add_heading(doc, "QUINTA. OBLIGACIONES DEL PRESTADOR", 1)
    for item in (
        "Prestar los Servicios con diligencia, personal competente y prácticas razonables de ingeniería y seguridad.",
        "Mantener reportes de avance y comunicar bloqueos, dependencias y riesgos relevantes.",
        "Usar los accesos de BEGLOBAL sólo para el alcance autorizado y aplicar privilegio mínimo cuando sea técnicamente posible.",
        "Corregir, dentro del periodo de garantía, defectos reproducibles atribuibles al alcance contratado.",
        "No publicar, comprar, transferir fondos, modificar cuentas críticas ni enviar campañas masivas sin aprobación humana expresa cuando tales acciones tengan efectos externos.",
        "Devolver o eliminar credenciales y materiales al concluir, conforme al Anexo de Datos y Seguridad.",
    ):
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "SEXTA. OBLIGACIONES DE BEGLOBAL", 1)
    for item in (
        "Designar un responsable con facultades para priorizar, aprobar y aceptar Entregables.",
        "Entregar oportunamente materiales, accesos, lineamientos de marca, metodología, respuestas y autorizaciones.",
        "Garantizar que los Materiales de BEGLOBAL y las instrucciones de uso sean lícitos y no vulneren derechos de terceros.",
        "Obtener avisos de privacidad, consentimientos y autorizaciones de Clientes Finales cuando corresponda.",
        "Revisar los Entregables y las salidas de IA antes de publicarlos o utilizarlos para decisiones, campañas, asesoría, transacciones o comunicaciones externas.",
        "Pagar honorarios, impuestos y costos autorizados en las fechas pactadas.",
        "Evitar el ingreso de contraseñas, datos bancarios, documentos oficiales, datos sensibles o información de menores en el Agente, salvo anexo específico y controles aprobados por ambas PARTES.",
    ):
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "SÉPTIMA. ENTREGA, ACEPTACIÓN Y ACTIVACIÓN", 1)
    add_body(
        doc,
        "El PRESTADOR notificará por escrito la disponibilidad de cada Entregable. BEGLOBAL contará con cinco (5) días hábiles para: (i) aceptarlo; o (ii) comunicar observaciones específicas, reproducibles y vinculadas con los criterios pactados. El PRESTADOR corregirá las no conformidades procedentes y presentará nuevamente el Entregable."
    )
    add_body(
        doc,
        "El Entregable se considerará aceptado si BEGLOBAL confirma su aceptación, lo pone en uso productivo, solicita la Activación o no presenta observaciones dentro del plazo. El silencio no subsanará incumplimientos ocultos ni limitará la garantía correctiva. Cada Agente podrá documentarse mediante el Acta de Activación del Anexo C."
    )

    add_heading(doc, "OCTAVA. HONORARIOS, IMPUESTOS Y PAGOS", 1)
    add_body(
        doc,
        "Para el piloto inicial descrito en el Anexo A, las PARTES acuerdan la contratación de tres (3) Agentes IA Ecommerce, a un precio de $2,500.00 MXN (dos mil quinientos pesos 00/100 M.N.) por cada Agente, IVA incluido, para un monto total de $7,500.00 MXN (siete mil quinientos pesos 00/100 M.N.), IVA incluido. El precio cubre configuración, personalización inicial, QA, Activación y treinta (30) días de soporte correctivo por Agente."
    )
    add_body(
        doc,
        "El PRESTADOR emitirá el CFDI correspondiente. Salvo pacto distinto, las facturas serán pagaderas dentro de cinco (5) días hábiles posteriores a su recepción, mediante transferencia a la cuenta comunicada por un canal oficial. Ningún cambio de cuenta será válido sin verificación adicional con el contacto autorizado."
    )
    add_body(
        doc,
        "Licencias, consumo de modelos, hosting, mensajería, dominios, APIs, tiendas, almacenamiento y demás cargos de terceros no están incluidos salvo indicación expresa. El PRESTADOR no contratará cargos extraordinarios por cuenta de BEGLOBAL sin autorización escrita. El incumplimiento de pago faculta al PRESTADOR a suspender Servicios no críticos después de aviso y un plazo de cinco (5) días hábiles para subsanar."
    )

    add_heading(doc, "NOVENA. PLAZO Y VIGENCIA", 1)
    add_body(
        doc,
        "El Contrato entrará en vigor en la fecha de su firma y permanecerá vigente por doce (12) meses, prorrogable por escrito. La Orden de Trabajo inicial tendrá una duración estimada de treinta (30) a cuarenta y cinco (45) días naturales, condicionada a la entrega oportuna de insumos y aprobaciones. La terminación del Contrato no extinguirá las Órdenes de Trabajo activas salvo acuerdo o causa de terminación."
    )

    add_heading(doc, "DÉCIMA. GARANTÍA Y SOPORTE", 1)
    add_body(
        doc,
        "Salvo pacto distinto, cada Agente contará con treinta (30) días naturales de soporte correctivo a partir de su Activación. La garantía cubre defectos reproducibles que impidan el funcionamiento conforme al alcance; no cubre nuevas funcionalidades, cambios de metodología, contenido adicional, errores de terceros, mal uso, modificaciones no autorizadas ni indisponibilidad externa."
    )
    add_body(
        doc,
        "Los tiempos de respuesta son objetivos de atención y no garantías de resolución. El soporte recurrente posterior, mantenimiento evolutivo, reentrenamiento, ampliaciones o acompañamiento adicional requerirán una Orden de Trabajo o plan de soporte."
    )

    add_heading(doc, "DÉCIMA PRIMERA. INTELIGENCIA ARTIFICIAL Y SUPERVISIÓN HUMANA", 1)
    for item in (
        "Las salidas de IA pueden ser inexactas, incompletas, sesgadas, no exclusivas o variar entre ejecuciones.",
        "BEGLOBAL conservará control humano sobre publicaciones, ofertas, precios, declaraciones legales o fiscales, decisiones de crédito, salud, empleo, pagos, compras y demás acciones de impacto material.",
        "El Agente no deberá presentarse como abogado, contador, médico, autoridad, humano real ni garantía de resultados.",
        "No se utilizará la imagen, voz o identidad de una persona para generar contenido sintético sin autorización expresa y documentada.",
        "El PRESTADOR no utilizará información confidencial de BEGLOBAL para entrenar modelos propios de propósito general. Los Servicios de Terceros se sujetarán a las opciones de privacidad y retención disponibles y aprobadas en el Anexo B.",
        "BEGLOBAL comunicará a los Clientes Finales, cuando resulte aplicable, que interactúan con un sistema automatizado y habilitará un mecanismo razonable de escalamiento humano.",
    ):
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "DÉCIMA SEGUNDA. DATOS PERSONALES Y SEGURIDAD", 1)
    add_body(
        doc,
        "Cuando el PRESTADOR trate datos personales por cuenta de BEGLOBAL, BEGLOBAL actuará como responsable y el PRESTADOR como persona encargada. El PRESTADOR tratará los datos únicamente conforme a instrucciones documentadas, para las finalidades del Servicio y durante el tiempo necesario."
    )
    for item in (
        "Aplicar medidas administrativas, técnicas y físicas razonables según el riesgo.",
        "Restringir accesos a personal y subencargados que necesiten conocer la información y estén sujetos a confidencialidad.",
        "Notificar a BEGLOBAL sin demora indebida y, de ser posible, dentro de veinticuatro (24) horas de conocer un incidente que pueda comprometer datos tratados por cuenta de BEGLOBAL.",
        "Colaborar razonablemente con BEGLOBAL en solicitudes de derechos ARCO, investigaciones y medidas de contención.",
        "Devolver o eliminar los datos al concluir el Servicio, salvo obligación legal de conservación.",
    ):
        add_list_item(doc, item, bullet_num)
    add_body(
        doc,
        "BEGLOBAL será responsable del aviso de privacidad, la base jurídica, la proporcionalidad de los datos y las instrucciones entregadas. Las categorías autorizadas, retención, transferencias y subencargados se detallan en el Anexo B."
    )

    add_heading(doc, "DÉCIMA TERCERA. CONFIDENCIALIDAD", 1)
    add_body(
        doc,
        "Cada PARTE protegerá la información técnica, comercial, financiera, estratégica, de seguridad, clientes, metodologías, accesos y demás información identificada como confidencial o que razonablemente deba entenderse como tal. Sólo podrá usarla para ejecutar el Contrato y revelarla a personal o asesores con necesidad de conocer y obligaciones equivalentes."
    )
    add_body(
        doc,
        "No será confidencial la información que la PARTE receptora demuestre que: (i) era pública sin infracción; (ii) ya conocía legítimamente; (iii) recibió lícitamente de un tercero; o (iv) desarrolló de forma independiente. Si una autoridad exige su revelación, la receptora notificará previamente cuando la ley lo permita."
    )
    add_body(
        doc,
        "La obligación subsistirá cinco (5) años después de la terminación; respecto de secretos industriales, credenciales y datos personales, subsistirá mientras conserven tal carácter o durante el plazo exigido por ley."
    )

    add_heading(doc, "DÉCIMA CUARTA. PROPIEDAD INTELECTUAL", 1)
    add_body(doc, "La titularidad se distribuirá de la siguiente forma:")
    for item in (
        "Materiales de BEGLOBAL. Permanecen bajo titularidad de BEGLOBAL o sus licenciantes. BEGLOBAL concede al PRESTADOR una licencia limitada, temporal y revocable para ejecutar los Servicios.",
        "Tecnología preexistente del PRESTADOR. Herramientas, plantillas, librerías, conectores, métodos, know-how, componentes reutilizables y mejoras generales preexistentes o independientes permanecen bajo titularidad del PRESTADOR.",
        "Entregables personalizados. Una vez pagados íntegramente, los derechos patrimoniales que legalmente puedan transmitirse sobre configuraciones, textos, estructuras y desarrollos creados específicamente por encargo para BEGLOBAL corresponderán a BEGLOBAL, con la amplitud necesaria para su explotación, salvo componentes preexistentes, de terceros o de código abierto.",
        "Licencia de componentes incorporados. El PRESTADOR concede a BEGLOBAL una licencia no exclusiva, mundial, durante la vigencia de los derechos y sin regalías adicionales para usar los componentes preexistentes incorporados que sean necesarios para aprovechar el Entregable, sin transferir herramientas generales separables.",
        "Contenido generado por IA. Las PARTES reconocen que la protección, exclusividad o titularidad de salidas generadas por IA puede depender de la intervención humana y la legislación aplicable. No se garantiza que una salida sea única o registrable.",
        "Créditos y portafolio. El PRESTADOR no usará nombre, marca, capturas, métricas ni Entregables de BEGLOBAL como caso de éxito sin autorización escrita.",
    ):
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "DÉCIMA QUINTA. DECLARACIONES Y GARANTÍAS", 1)
    add_body(
        doc,
        "Cada PARTE declara que tiene derecho a celebrar el Contrato. BEGLOBAL garantiza que los materiales e instrucciones que proporcione cuentan con las licencias, permisos y consentimientos necesarios. El PRESTADOR garantiza que no incorporará deliberadamente código malicioso y que respetará las licencias de terceros que identifique en los Entregables."
    )
    add_body(
        doc,
        "Si una reclamación de tercero se deriva de materiales o instrucciones de una PARTE, dicha PARTE asumirá la defensa y los daños definitivamente impuestos o acordados, siempre que reciba aviso oportuno, control razonable de la defensa y cooperación. No podrán celebrarse acuerdos que admitan responsabilidad o impongan obligaciones a la otra PARTE sin su consentimiento."
    )

    add_heading(doc, "DÉCIMA SEXTA. RESPONSABILIDAD", 1)
    add_body(
        doc,
        "Cada PARTE responderá por los daños directos que cause por incumplimiento probado. En la medida permitida por la ley, ninguna PARTE será responsable por daños indirectos, lucro cesante, pérdida de oportunidad, pérdida de datos no respaldados o decisiones tomadas exclusivamente con base en salidas de IA."
    )
    add_body(
        doc,
        "Salvo dolo, negligencia grave, violación de confidencialidad, infracción de propiedad intelectual, incumplimiento de obligaciones de datos personales, obligaciones de indemnización o falta de pago, la responsabilidad acumulada de cada PARTE derivada de una Orden de Trabajo no excederá los honorarios efectivamente pagados o pagaderos bajo dicha Orden durante los seis (6) meses anteriores al hecho. Nada limita responsabilidades que legalmente no puedan limitarse."
    )

    add_heading(doc, "DÉCIMA SÉPTIMA. INDEPENDENCIA Y AUSENCIA DE RELACIÓN LABORAL", 1)
    add_body(
        doc,
        "El PRESTADOR actuará con autonomía técnica, administrativa y organizativa, sin exclusividad, subordinación, horario impuesto ni integración a la estructura laboral de BEGLOBAL. Cada PARTE será responsable de su personal, impuestos, seguridad social, herramientas y dirección. Ninguna estipulación prevalecerá sobre la realidad de la prestación ni impedirá la aplicación de normas laborales imperativas si existiera subordinación efectiva."
    )

    add_heading(doc, "DÉCIMA OCTAVA. SUSPENSIÓN Y TERMINACIÓN", 1)
    add_body(
        doc,
        "Cualquiera de las PARTES podrá terminar el Contrato u Orden de Trabajo por incumplimiento material no subsanado dentro de diez (10) días hábiles después de una notificación detallada. También podrá terminar por insolvencia, uso ilícito, riesgo grave de seguridad o instrucción contraria a derecho, en cuyo caso la suspensión podrá ser inmediata y deberá documentarse."
    )
    add_body(
        doc,
        "Cualquiera de las PARTES podrá terminar por conveniencia con quince (15) días naturales de aviso. BEGLOBAL pagará los Servicios efectivamente prestados, Activaciones realizadas y compromisos de terceros previamente autorizados y no cancelables. El PRESTADOR entregará los avances pagados y prestará asistencia de transición razonable cotizada por separado."
    )
    add_body(
        doc,
        "A la terminación sobrevivirán las obligaciones de pago, confidencialidad, datos, propiedad intelectual, responsabilidad, solución de controversias y cualquier otra que por su naturaleza deba subsistir."
    )

    add_heading(doc, "DÉCIMA NOVENA. FUERZA MAYOR", 1)
    add_body(
        doc,
        "Ninguna PARTE será responsable por retrasos derivados de acontecimientos fuera de su control razonable, incluidos desastres, fallas generalizadas de internet o energía, actos de autoridad, conflictos, epidemias o indisponibilidad extraordinaria de proveedores esenciales. La PARTE afectada notificará, mitigará y reanudará el cumplimiento tan pronto sea razonable. Las obligaciones de pago ya vencidas no se suspenden."
    )

    add_heading(doc, "VIGÉSIMA. COMUNICACIONES Y APROBACIONES", 1)
    add_body(
        doc,
        "Las notificaciones formales se enviarán a los correos señalados en el bloque de firmas o en la Orden de Trabajo y se considerarán recibidas al existir acuse, respuesta o evidencia técnica de entrega sin rebote. WhatsApp, Telegram, Slack u otros canales podrán utilizarse para coordinación y aprobaciones operativas cuando el mensaje provenga de un contacto autorizado, pero no para modificar responsabilidad, propiedad intelectual, precio o jurisdicción sin confirmación formal."
    )

    add_heading(doc, "VIGÉSIMA PRIMERA. FIRMA ELECTRÓNICA Y EVIDENCIA", 1)
    add_body(
        doc,
        "Las PARTES podrán firmar autógrafa o electrónicamente y aceptar Órdenes de Trabajo mediante plataformas de firma o mensajes de datos atribuibles, íntegros y accesibles para consulta posterior. Las copias y contrapartes electrónicas producirán los mismos efectos entre las PARTES, sin perjuicio de los requisitos imperativos aplicables."
    )

    add_heading(doc, "VIGÉSIMA SEGUNDA. CESIÓN Y SUBCONTRATACIÓN", 1)
    add_body(
        doc,
        "Ninguna PARTE cederá el Contrato sin consentimiento escrito de la otra, salvo reorganización corporativa que no reduzca garantías. El PRESTADOR podrá utilizar personal y subcontratistas bajo su responsabilidad. Los subencargados que traten datos personales deberán estar autorizados conforme al Anexo B y sujetos a obligaciones equivalentes."
    )

    add_heading(doc, "VIGÉSIMA TERCERA. LEY APLICABLE Y CONTROVERSIAS", 1)
    p = doc.add_paragraph(style="Clause Body")
    r = p.add_run("El Contrato se regirá por las leyes de los Estados Unidos Mexicanos. Antes de iniciar un procedimiento, las PARTES intentarán resolver la controversia mediante negociación entre responsables designados durante diez (10) días hábiles. Si no hay acuerdo, se someten a los tribunales competentes de ")
    set_run_font(r)
    r = p.add_run("Hermosillo, Sonora")
    set_run_font(r)
    r = p.add_run(", renunciando al fuero que pudiera corresponderles por domicilios presentes o futuros, salvo competencia irrenunciable.")
    set_run_font(r)

    add_heading(doc, "VIGÉSIMA CUARTA. DISPOSICIONES GENERALES", 1)
    for item in (
        "Integridad. El Contrato y sus anexos sustituyen entendimientos previos sobre su objeto.",
        "Modificaciones. Sólo serán válidas por escrito y aceptadas por representantes autorizados.",
        "Divisibilidad. La invalidez de una disposición no afectará las restantes; las PARTES la sustituirán por una válida de efecto equivalente.",
        "No renuncia. La falta de ejercicio de un derecho no implica renuncia.",
        "Encabezados. Se utilizan para referencia y no limitan la interpretación.",
        "Anticorrupción. Las PARTES cumplirán la legislación aplicable y no ofrecerán beneficios indebidos relacionados con el Contrato.",
    ):
        add_list_item(doc, item, bullet_num)

    add_body(
        doc,
        "Leído el presente Contrato y enteradas las PARTES de su contenido y alcance, lo firman en dos ejemplares o mediante firma electrónica en la fecha indicada."
    )

    add_signature_table(
        doc,
        [
            ("SISTEMA BE GLOBAL PRO", "[NOMBRE COMPLETO DEL REPRESENTANTE LEGAL]"),
            ("EL PRESTADOR / SOFTVIBES", "[JOSÉ ROGELIO GARCÍA VITAL / REPRESENTANTE]"),
        ],
    )

    p = doc.add_paragraph(style="Small Note")
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Correos para notificaciones: BEGLOBAL [CORREO] | PRESTADOR [CORREO]")
    set_run_font(r, size=9, color=MUTED)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.add_page_break()
    add_heading(doc, "ANEXO A. ORDEN DE TRABAJO INICIAL", 1)
    add_body(doc, "Piloto de Agentes IA Ecommerce para BeGlobal", italic=True)

    add_kv_table(
        doc,
        [
            ("Proyecto", "Piloto Agente IA Ecommerce BeGlobal", False),
            ("Objetivo", "Validar una experiencia guiada de ecommerce con usuarios reales y obtener evidencia antes de escalar.", False),
            ("Duración estimada", "30 a 45 días naturales desde la reunión de arranque y entrega de insumos.", False),
            ("Volumen contratado", "3 Agentes IA Ecommerce", False),
            ("Responsable BEGLOBAL", "[NOMBRE, CARGO Y CORREO]", True),
            ("Responsable PRESTADOR", "[NOMBRE, CARGO Y CORREO]", True),
            ("Canal operativo", "[WHATSAPP / TELEGRAM / SLACK / CORREO]", True),
        ],
        widths=(2700, 6660),
    )

    add_heading(doc, "A.1 Alcance incluido", 2)
    for item in (
        "Sesión de descubrimiento y levantamiento del flujo BeGlobal.",
        "Configuración de un Agente IA Ecommerce base bajo marca BeGlobal.",
        "Personalización por Cliente Final con nombre, negocio, etapa y necesidades iniciales.",
        "Carga inicial de metodología, preguntas frecuentes, contenidos y guardrails aprobados por BEGLOBAL.",
        "Diagnóstico, siguiente paso, ideas de contenido, títulos, descripciones, ofertas, scripts básicos y tareas semanales.",
        "Pruebas funcionales, ajuste inicial, Activación y documentación de uso.",
        "Capacitación breve al responsable operativo de BEGLOBAL.",
        "Soporte correctivo durante 30 días naturales por Agente desde su Activación.",
        "Reporte final del piloto con uso, incidencias, preguntas frecuentes y recomendaciones, sujeto a disponibilidad de datos.",
    ):
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "A.2 Entregables y aceptación", 2)
    add_data_table(
        doc,
        ["#", "Entregable", "Evidencia", "Criterio de aceptación"],
        [
            ("1", "Brief y mapa de flujo", "Documento aprobado", "Fases, responsables y límites definidos"),
            ("2", "Agente base BeGlobal", "Acceso de prueba", "Responde dentro del alcance y muestra guardrails"),
            ("3", "Personalización por usuario", "Ficha/registro", "Identidad y contexto mínimo configurados"),
            ("4", "Base inicial de conocimiento", "Inventario de fuentes", "Fuentes autorizadas y trazables"),
            ("5", "Pruebas y Activación", "Acta del Anexo C", "Casos críticos superados o excepciones aceptadas"),
            ("6", "Capacitación y guía", "Sesión + documento", "Responsable conoce uso, soporte y escalamiento"),
            ("7", "Reporte de cierre", "Reporte ejecutivo", "Hallazgos y recomendación de siguiente fase"),
        ],
        [540, 2340, 2160, 4320],
        sizes=[8.5, 8.8, 8.8, 8.8],
    )

    add_heading(doc, "A.3 Calendario de referencia", 2)
    add_data_table(
        doc,
        ["Fase", "Periodo", "Actividad principal", "Dependencia"],
        [
            ("1", "Semana 1", "Alineación, materiales, riesgos y criterios", "Responsables e insumos"),
            ("2", "Semana 2", "Configuración del Agente base y contenido", "Aprobación metodológica"),
            ("3", "Semana 3", "Personalización, QA y primeras Activaciones", "Usuarios piloto"),
            ("4", "Semanas 4-6", "Uso controlado, ajustes y reporte", "Participación y evidencia"),
        ],
        [900, 1440, 3960, 3060],
        sizes=[9, 9, 9, 9],
    )

    add_heading(doc, "A.4 Condiciones comerciales propuestas", 2)
    add_kv_table(
        doc,
        [
            ("Tarifa por Agente", "$2,500.00 MXN por Agente, IVA incluido", False),
            ("Monto total contratado", "$7,500.00 MXN por 3 Agentes, IVA incluido", False),
            ("Naturaleza del cargo", "Pago único por configuración, personalización inicial, QA, Activación y 30 días de soporte correctivo.", False),
            ("Facturación", "[POR ACTIVACIÓN / POR LOTE / ANTICIPO Y SALDO]", True),
            ("Vencimiento", "5 días hábiles desde recepción de CFDI, salvo acuerdo distinto.", False),
            ("Costos recurrentes", "[DEFINIR HOSTING, MODELOS, MENSAJERÍA, ALMACENAMIENTO Y SOPORTE POSTERIOR]", True),
            ("Moneda e impuestos", "Pesos mexicanos; los importes anteriores incluyen IVA. El CFDI desglosará los impuestos conforme a ley.", False),
        ],
        widths=(2700, 6660),
    )

    add_heading(doc, "A.5 Exclusiones", 2)
    for item in (
        "Integraciones avanzadas con CRM, ERP, marketplaces, pagos, bancos o cuentas publicitarias.",
        "Automatizaciones con efectos externos sin aprobación humana.",
        "Carga masiva de catálogos, migraciones o limpieza extensiva de datos.",
        "Administración completa del negocio, soporte directo ilimitado o acompañamiento uno a uno permanente.",
        "Asesoría legal, fiscal, contable, financiera o médica.",
        "Garantías de ventas, conversiones, ingresos, disponibilidad o aprobación por plataformas.",
        "Desarrollos a medida no descritos y cargos de Servicios de Terceros.",
    ):
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "A.6 Métricas sugeridas", 2)
    for item in (
        "Usuarios activados y usuarios que completan el diagnóstico.",
        "Conversaciones y tareas guiadas por usuario.",
        "Preguntas frecuentes, bloqueos y escalamientos humanos.",
        "Entregables generados y aceptados por los usuarios.",
        "Incidencias, tiempos de atención y causas de abandono.",
        "Percepción de utilidad y recomendación de escalamiento.",
    ):
        add_list_item(doc, item, bullet_num)

    add_signature_table(
        doc,
        [
            ("BEGLOBAL - APROBACIÓN DE LA ORDEN", "[NOMBRE]"),
            ("PRESTADOR - APROBACIÓN DE LA ORDEN", "[NOMBRE]"),
        ],
    )

    doc.add_page_break()
    add_heading(doc, "ANEXO B. DATOS PERSONALES, SEGURIDAD Y TERCEROS", 1)

    add_heading(doc, "B.1 Roles e instrucciones", 2)
    add_body(
        doc,
        "BEGLOBAL determina las finalidades y medios esenciales del tratamiento y actúa como responsable. El PRESTADOR actúa como persona encargada únicamente respecto de los datos tratados por cuenta de BEGLOBAL. El PRESTADOR no venderá, perfilará para fines propios ni utilizará esos datos para entrenar modelos de propósito general."
    )

    add_heading(doc, "B.2 Categorías autorizadas", 2)
    add_data_table(
        doc,
        ["Categoría", "Ejemplos", "Autorización"],
        [
            ("Identificación básica", "Nombre, alias, empresa, rol", "Permitida si es necesaria"),
            ("Contacto", "Correo, teléfono o usuario de mensajería", "Permitida con aviso aplicable"),
            ("Contexto de negocio", "Etapa, productos, canales, objetivos", "Permitida"),
            ("Contenido", "Textos, imágenes, videos y documentos", "Sólo con derechos y finalidad válida"),
            ("Datos sensibles", "Salud, biométricos, creencias, origen, orientación", "Prohibida salvo anexo específico"),
            ("Financieros/credenciales", "Tarjetas, banca, contraseñas, tokens", "Prohibida en conversaciones; usar canales seguros"),
            ("Menores", "Datos de personas menores de edad", "Prohibida salvo evaluación y autorización jurídica"),
        ],
        [2160, 4140, 3060],
        sizes=[8.8, 8.8, 8.8],
    )

    add_heading(doc, "B.3 Medidas mínimas", 2)
    for item in (
        "Control de acceso individual, privilegio mínimo y revocación oportuna.",
        "Cifrado en tránsito y, cuando esté disponible, en reposo.",
        "Separación razonable de entornos y datos por cliente o perfil.",
        "Registro de cambios y eventos relevantes según capacidades de la plataforma.",
        "Respaldo y recuperación para información crítica expresamente contratada.",
        "Gestión de vulnerabilidades y actualizaciones razonables.",
        "No insertar secretos en código, repositorios públicos, prompts compartidos o canales no autorizados.",
        "Notificación, contención, análisis y cooperación ante incidentes.",
    ):
        add_list_item(doc, item, bullet_num)

    add_heading(doc, "B.4 Proveedores y subencargados autorizados", 2)
    add_body(
        doc,
        "Antes de producción, las PARTES completarán esta tabla. Todo proveedor adicional que trate datos personales requerirá aviso y autorización conforme al Contrato."
    )
    add_data_table(
        doc,
        ["Proveedor", "Servicio", "Datos", "Región/transferencia", "Retención"],
        [
            ("[PROVEEDOR DE IA]", "Modelo de lenguaje", "[CATEGORÍAS]", "[PAÍS/REGIÓN]", "[PLAZO]"),
            ("[HOSTING]", "Infraestructura", "[CATEGORÍAS]", "[PAÍS/REGIÓN]", "[PLAZO]"),
            ("[MENSAJERÍA]", "Canal de usuario", "[CATEGORÍAS]", "[PAÍS/REGIÓN]", "[PLAZO]"),
        ],
        [1800, 1800, 1980, 2160, 1620],
        sizes=[8.2, 8.2, 8.2, 8.2, 8.2],
    )

    add_heading(doc, "B.5 Retención, devolución y eliminación", 2)
    add_body(
        doc,
        "Salvo obligación legal o instrucción documentada, el PRESTADOR devolverá o eliminará los datos y credenciales dentro de quince (15) días hábiles posteriores a la terminación o solicitud de BEGLOBAL. Los respaldos se eliminarán conforme al ciclo técnico aplicable y permanecerán protegidos mientras existan. BEGLOBAL conservará sus propios respaldos de Materiales y Entregables."
    )

    add_heading(doc, "B.6 Incidentes", 2)
    add_body(
        doc,
        "La notificación preliminar incluirá, cuando esté disponible: naturaleza del incidente, fecha estimada, sistemas y datos afectados, medidas adoptadas, riesgos y contacto de seguimiento. Ninguna PARTE notificará públicamente o a titulares en nombre de la otra sin coordinación, salvo obligación legal inmediata."
    )

    add_signature_table(
        doc,
        [
            ("BEGLOBAL - RESPONSABLE", "[NOMBRE]"),
            ("PRESTADOR - PERSONA ENCARGADA", "[NOMBRE]"),
        ],
    )

    doc.add_page_break()
    add_heading(doc, "ANEXO C. ACTA DE ACTIVACIÓN Y ACEPTACIÓN", 1)
    add_body(
        doc,
        "Utilizar una copia por Agente o lote. La firma confirma la Activación técnica; no modifica el alcance, precio ni responsabilidades del Contrato."
    )
    add_kv_table(
        doc,
        [
            ("Orden de Trabajo", "[IDENTIFICADOR]", True),
            ("Cliente Final / lote", "[NOMBRE O CÓDIGO]", True),
            ("Agente / perfil", "[NOMBRE E IDENTIFICADOR]", True),
            ("Canal", "[TELEGRAM / WEB / OTRO]", True),
            ("Fecha de Activación", "[FECHA Y HORA]", True),
            ("Inicio de garantía", "[FECHA]", True),
            ("Fin de garantía", "[FECHA]", True),
            ("Responsable BEGLOBAL", "[NOMBRE Y CORREO]", True),
            ("Responsable PRESTADOR", "[NOMBRE Y CORREO]", True),
        ],
        widths=(2700, 6660),
    )

    add_heading(doc, "C.1 Verificaciones", 2)
    checks = (
        "[ ] Identidad y contexto mínimo configurados.",
        "[ ] Fuentes y metodología autorizadas cargadas o vinculadas.",
        "[ ] Casos de prueba críticos ejecutados.",
        "[ ] Guardrails, privacidad y escalamiento humano revisados.",
        "[ ] Accesos de producción entregados por canal seguro.",
        "[ ] Costos recurrentes y proveedores confirmados.",
        "[ ] Observaciones abiertas registradas y aceptadas.",
    )
    for check in checks:
        p = doc.add_paragraph(style="Clause Body")
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(check)
        set_run_font(r)

    add_heading(doc, "C.2 Resultado", 2)
    add_kv_table(
        doc,
        [
            ("Estado", "[ ] Aceptado  [ ] Aceptado con observaciones  [ ] Rechazado con causa", True),
            ("Observaciones", "[DETALLAR O REFERENCIAR TICKET]", True),
            ("Fecha límite de corrección", "[FECHA / NO APLICA]", True),
        ],
        widths=(2700, 6660),
    )
    add_signature_table(
        doc,
        [
            ("BEGLOBAL - ACEPTACIÓN", "[NOMBRE]"),
            ("PRESTADOR - ENTREGA", "[NOMBRE]"),
        ],
    )

    configure_page(doc)
    for section in doc.sections:
        set_header_footer(section)

    doc.core_properties.title = "Contrato marco de prestación de servicios - Sistema Be Global Pro y Softvibes"
    doc.core_properties.subject = "Agentes IA Ecommerce"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.keywords = "Sistema Be Global Pro, Softvibes, prestación de servicios, agentes IA"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
